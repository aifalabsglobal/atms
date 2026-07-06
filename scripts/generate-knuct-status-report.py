"""Generate KNUCT Integration Status Report as Word document."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = '1A3C6E') -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = None
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = text
            for p in table.rows[r_idx].cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def main() -> None:
    out = Path(__file__).resolve().parents[1] / 'docs' / 'KNUCT_Integration_Status_Report.docx'
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('KNUCT Blockchain Integration\nStatus Report')
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('JNTUH Academic Tracking & Management System (ATMS)').font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Report Date: {date.today().strftime("%d %B %Y")}\nVersion: 0.2.0').font.size = Pt(10)

    doc.add_paragraph()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report summarises the current status of integrating Knuct Blockchain DID Authentication '
        'into the JNTUH ATMS platform. Core DID authentication has been implemented based on the '
        'KnuctBlockchainDID-Auth-API-v2 specification and the did-auth-modules package provided by Knuct. '
        'Wallet provisioning, encrypted private share storage, private share download, and a user-facing '
        'DID authentication panel are operational in the development environment. End-to-end validation '
        'with vendor-supplied test credentials, CAPI dashboard integration, mobile support, and '
        'production-scale session handling remain pending.'
    )

    # 2. Vendor Materials
    doc.add_heading('2. Vendor Materials Received', level=1)
    add_table(doc,
        ['Material', 'Location / Description'],
        [
            ['KnuctBlockchainDID-Auth-API-v2.pdf', 'D:\\Yoga_Data\\Atlas\\ — DID Auth API specification'],
            ['did-auth-modules.zip', 'D:\\Yoga_Data\\Atlas\\did-auth-modules\\ — Reference React modules'],
            ['nlss.js', 'NLSS cryptographic challenge-response algorithm'],
            ['privShare.js', 'Private share image processing utilities'],
            ['ServerAPI.js', 'SAPI wrappers (/sapi/* endpoints)'],
            ['ClientAPI.js', 'CAPI wrappers (/capi/* endpoints)'],
            ['AccessWallet.js', 'Reference DID auth React flow'],
        ]
    )

    # 3. Completed Work
    doc.add_heading('3. Completed Work', level=1)

    doc.add_heading('3.1 Dependencies Installed', level=2)
    add_table(doc, ['Package', 'Version', 'Purpose'], [
        ['js-sha3', '0.9.3', 'SHA3-256 hashing for NLSS algorithm'],
        ['md5', '2.3.0', 'Multihash MD5 for private share image hashing'],
        ['@types/md5', 'dev', 'TypeScript type definitions'],
    ])

    doc.add_heading('3.2 New / Modified Application Files', level=2)
    add_table(doc, ['File', 'Description'], [
        ['src/lib/knuct/nlss.ts', 'Port of nlss.js — NLSS challenge-response signing'],
        ['src/lib/knuct/priv-share.ts', 'Port of privShare.js — image hash pipeline'],
        ['src/lib/knuct/did-auth-session.ts', 'In-memory server session store for DID auth'],
        ['src/lib/knuct/knuct-client.ts', 'Added authChallenge, authResponse, startNode, walletData, logout'],
        ['src/lib/knuct/mock-adapter.ts', 'Mock implementations for dev/testing'],
        ['src/lib/knuct/types.ts', 'Extended KnuctAdapter interface'],
        ['src/app/api/knuct/did-auth/route.ts', 'Two-step DID auth API (challenge + complete)'],
        ['src/app/api/knuct/privshare/route.ts', 'Authenticated private share download'],
        ['src/components/knuct/did-auth-panel.tsx', 'Client-side DID auth UI component'],
        ['src/components/sections/settings-section.tsx', 'Knuct tab — DID panel + download button'],
    ])

    doc.add_heading('3.3 DID Authentication Flow (Implemented)', level=2)
    add_table(doc, ['Step', 'Knuct Endpoint', 'Implementation Status'], [
        ['1', 'POST /sapi/auth/challenge', 'Complete — hash sent, challenge returned'],
        ['2', 'POST /sapi/auth/response', 'Complete — NLSS response array submitted'],
        ['3', 'GET /sapi/startnode', 'Complete — IPFS/blockchain node started'],
        ['4', 'GET /sapi/walletdata', 'Complete — DID retrieved and saved to database'],
    ])

    doc.add_heading('3.4 Wallet & Private Share', level=2)
    for item in [
        'Wallet provisioning via POST /sapi/createwallet (existing, enhanced)',
        'Private share fetched from Knuct and stored AES-256-GCM encrypted in PostgreSQL',
        'Download endpoint: GET /api/knuct/privshare (authenticated users only)',
        'Settings UI: Provision wallet, Download private share, DID Authentication panel',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.5 Pre-existing Knuct Integration (ATMS)', level=2)
    for item in [
        'Hash anchors / audit trail in PostgreSQL',
        'Circuit breaker and health checks (Knuct ping on /api/health)',
        'Pilot cohort support (KNUCT_PILOT_COHORT_LIMIT)',
        'Mock adapter for development when Knuct unavailable',
        'Knuct configuration via environment variables',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # 4. Architecture
    doc.add_heading('4. Integration Architecture', level=1)
    doc.add_paragraph(
        'Knuct documentation states that DID authentication should run at the front-end level. '
        'ATMS uses a hybrid architecture:'
    )
    add_table(doc, ['Operation', 'Execution Location'], [
        ['Private share image read & pixel processing', 'Browser (client-side)'],
        ['Hash computation (md5 multihash + base32)', 'Browser (client-side)'],
        ['NLSS challenge-response computation', 'Browser (client-side)'],
        ['Calls to /sapi/auth/challenge, /response, /startnode, /walletdata', 'ATMS Next.js server (proxy)'],
        ['Knuct session cookie management', 'ATMS server (in-memory, per user, 10-min TTL)'],
        ['DID persistence', 'ATMS PostgreSQL database'],
    ])
    doc.add_paragraph(
        'Note: The private share image is never uploaded during DID authentication. Only the '
        'computed hash and binary response array are transmitted to the ATMS server.'
    )

    # 5. Environment
    doc.add_heading('5. Current Configuration', level=1)
    add_table(doc, ['Variable', 'Current Value / Status'], [
        ['KNUCT_ENABLED', 'true'],
        ['KNUCT_BASE_URL', 'https://webwallet.knuct.com'],
        ['KNUCT_PRIVSHARE_ENC_KEY', 'Configured (64-char hex)'],
        ['KNUCT_WALLET_ON_USER_CREATE', 'false'],
        ['KNUCT_PILOT_COHORT_LIMIT', '25'],
        ['KNUCT_API_KEY / KNUCT_API_SECRET', 'Not configured — pending vendor guidance'],
        ['Development server', 'http://localhost:3000'],
    ])

    # 6. Security
    doc.add_heading('6. Security Measures', level=1)
    for item in [
        'Private share encrypted at rest using AES-256-GCM (KNUCT_PRIVSHARE_ENC_KEY)',
        'During DID auth, only hash and response array sent — not the image file',
        'All Knuct API routes require authenticated ATMS session (NextAuth)',
        'Private share download restricted to owning user',
        'DID auth server sessions expire after 10 minutes',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # 7. Pending
    doc.add_heading('7. Pending Work', level=1)

    doc.add_heading('7.1 Not Yet Built', level=2)
    add_table(doc, ['Item', 'Details', 'Priority'], [
        ['CAPI integration', '/capi/getDashboard, getAccountInfo, start, check — wallet stats on dashboard', 'Medium'],
        ['Mobile private share', 'OpenCV (opencv.js) path for mobile browsers not integrated', 'Medium'],
        ['Logout UI', '/sapi/logout implemented in code but not exposed in UI', 'Low'],
        ['Session expiry UX', 'Knuct 15-min session — retry/re-auth flow not polished', 'Medium'],
        ['Production session store', 'In-memory sessions — needs Redis or sticky sessions for scale', 'High'],
    ])

    doc.add_heading('7.2 Built but Not Validated', level=2)
    add_table(doc, ['Item', 'Details'], [
        ['End-to-end DID auth', 'Code complete; not tested with vendor-supplied real private share'],
        ['Architecture sign-off', 'Hybrid model requires Knuct confirmation'],
        ['startnode vs starttempnode', 'Both used — vendor clarification needed'],
    ])

    doc.add_heading('7.3 Awaiting Knuct', level=2)
    add_table(doc, ['Requirement', 'Purpose'], [
        ['Test wallet + sample private share', 'Live end-to-end validation'],
        ['Architecture confirmation', 'Approve hybrid client/server model'],
        ['CAPI documentation', 'Dashboard wallet statistics integration'],
        ['Sandbox / test environment', 'Safe integration testing'],
        ['API key / tenant requirements', 'Production deployment configuration'],
        ['Mobile / OpenCV guidance', 'Phone browser support'],
        ['Rate limits', 'Especially on /sapi/auth/challenge'],
    ])

    # 8. CAPI
    doc.add_heading('8. CAPI Overview (Not Yet Integrated)', level=1)
    doc.add_paragraph(
        'CAPI (Client API) is Knuct\'s post-authentication wallet client layer at '
        'https://webwallet.knuct.com/capi. It is separate from SAPI (Server API) which handles '
        'authentication and wallet creation.'
    )
    add_table(doc, ['CAPI Endpoint', 'Purpose'], [
        ['GET /capi/check', 'Check if user is registered'],
        ['GET /capi/start', 'Initial IPFS setup and network node sync'],
        ['GET /capi/getAccountInfo', 'User account details'],
        ['GET /capi/getDashboard', 'Wallet statistics for dashboard display'],
    ])

    # 9. Next Steps
    doc.add_heading('9. Recommended Next Steps', level=1)
    steps = [
        'Send progress update email to Knuct partners (request test credentials and architecture confirmation).',
        'Run internal E2E test: Provision wallet → Download private share → Upload in DID Auth panel.',
        'Fix any API/flow mismatches identified during live testing.',
        'Implement Redis-backed DID auth sessions before production multi-instance deployment.',
        'Integrate CAPI endpoints once vendor documentation is received.',
        'Add mobile OpenCV support after vendor guidance.',
        'Deploy to staging and run pilot with cohort (up to 25 users per KNUCT_PILOT_COHORT_LIMIT).',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    # 10. Conclusion
    doc.add_heading('10. Conclusion', level=1)
    doc.add_paragraph(
        'The core Knuct DID authentication integration is substantially complete in the ATMS codebase. '
        'The application can provision wallets, securely store and download private shares, and '
        'authenticate users via the four-step DID flow described in Knuct\'s v2 API specification. '
        'Progress to production readiness depends on vendor-provided test credentials, architecture '
        'confirmation, CAPI integration for dashboard features, and production-grade session management.'
    )

    doc.add_paragraph()
    footer = doc.add_paragraph('— End of Report —')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out)
    print(f'Report saved to: {out}')


if __name__ == '__main__':
    main()
