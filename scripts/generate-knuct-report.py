"""Convert KNUCT_SCMS_Integration_Report.md to Word (.docx) in project root."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "KNUCT_SCMS_Integration_Report.md"
OUT_PATH = ROOT / "KNUCT_SCMS_Integration_Report.docx"

NAVY = RGBColor(0x1A, 0x3C, 0x6E)
CODE_FONT = "Consolas"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text: str) -> None:
    """Parse **bold** and *italic* inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(col_count):
            val = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, val.strip())
            if ri == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "E8EEF4")
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], lang: str = "") -> None:
    if lang == "mermaid":
        note = doc.add_paragraph()
        run = note.add_run("[Architecture diagram — see markdown source for mermaid rendering]")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(8)


def parse_markdown_table(lines: list[str]) -> list[list[str]] | None:
    if len(lines) < 2:
        return None
    if not all("|" in line for line in lines[:2]):
        return None
    sep = lines[1].strip()
    if not re.match(r"^\|?[\s\-:|]+\|?$", sep):
        return None

    def split_row(line: str) -> list[str]:
        line = line.strip().strip("|")
        return [c.strip() for c in line.split("|")]

    return [split_row(line) for line in lines if line.strip()]


def convert_md_to_docx(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    code_lang = ""
    list_buffer: list[str] = []

    def flush_list() -> None:
        nonlocal list_buffer
        for item in list_buffer:
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, item)
        list_buffer = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines, code_lang)
                code_lines = []
                code_lang = ""
                in_code = False
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_list()
            fence = line.strip()[3:].strip()
            code_lang = fence
            in_code = True
            i += 1
            continue

        if line.strip() == "---":
            flush_list()
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            flush_list()
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.color.rgb = NAVY
            i += 1
            continue

        if line.startswith("## "):
            flush_list()
            h = doc.add_heading(line[3:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = NAVY
            i += 1
            continue

        if line.startswith("### "):
            flush_list()
            h = doc.add_heading(line[4:].strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = NAVY
            i += 1
            continue

        if line.startswith("#### "):
            flush_list()
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.strip().startswith("|"):
            flush_list()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            parsed = parse_markdown_table(table_lines)
            if parsed:
                add_table(doc, parsed)
            continue

        if re.match(r"^[-*]\s+", line):
            flush_list()
            list_buffer.append(re.sub(r"^[-*]\s+", "", line.strip()))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            flush_list()
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, re.sub(r"^\d+\.\s+", "", line.strip()))
            i += 1
            continue

        if line.strip() == "":
            flush_list()
            i += 1
            continue

        if line.strip().startswith("*End of Report*"):
            flush_list()
            p = doc.add_paragraph("— End of Report —")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        flush_list()
        p = doc.add_paragraph()
        add_formatted_runs(p, line.strip())
        i += 1

    flush_list()


def main() -> None:
    if not MD_PATH.exists():
        raise SystemExit(f"Markdown source not found: {MD_PATH}")

    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    convert_md_to_docx(md_text, doc)
    doc.save(OUT_PATH)
    print(f"Written: {OUT_PATH} ({OUT_PATH.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
