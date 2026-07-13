"""Generate Knuct Integration Process & Future Requirements Word document (repo root)."""
from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "KNUCT_Integration_Process_and_Roadmap.docx"
NAVY = "1A3C6E"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = NAVY,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def git_tip() -> str:
    try:
        return (
            subprocess.check_output(
                ["git", "rev-parse", "--short", "HEAD"],
                cwd=ROOT,
                stderr=subprocess.DEVNULL,
                text=True,
            ).strip()
            or "unknown"
        )
    except Exception:
        return "unknown"


def main() -> None:
    tip = git_tip()
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Knuct Integration Process\n& Future Requirements")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "AIMSCS — Academic Tracking & Management System (ATMS)"
    ).font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"How Knuct was integrated end-to-end, and what is still required\n"
        f"Report date: {date.today().strftime('%d %B %Y')}  ·  Git tip: {tip}"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # ── 1. Purpose ───────────────────────────────────────────────────────
    doc.add_heading("1. Purpose of This Document", level=1)
    doc.add_paragraph(
        "This document describes the complete Knuct integration process as implemented "
        "in ATMS today, and lists what is still required for production-ready campus "
        "rollout. It complements KNUCT_Integration_Pointers.docx (file-level inventory)."
    )

    # ── 2. Goals ─────────────────────────────────────────────────────────
    doc.add_heading("2. Integration Goals", level=1)
    for b in [
        "Decentralized identity (DID) for students and staff via Knuct wallets.",
        "Optional DID login and self-registration alongside password auth.",
        "Tamper-evident audit trail: hash anchors for critical campus events.",
        "Future verifiable credentials (attendance certificates, transcripts).",
        "PostgreSQL remains the system of record; Knuct is an adapter behind SCMS APIs.",
        "Safe rollback: disable live Knuct without breaking core attendance/LMS.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # ── 3. Architecture principles ───────────────────────────────────────
    doc.add_heading("3. Architecture Principles", level=1)
    doc.add_paragraph(
        "Knuct is never called directly from UI business logic. All vendor I/O goes "
        "through src/lib/knuct/* (adapter pattern), similar to face-verification stubs."
    )
    add_table(
        doc,
        ["Layer", "Responsibility", "Key paths"],
        [
            ["UI", "Login/register DID panels, wallet, Settings → Knuct, /verify", "src/app/login, components/knuct"],
            ["API", "Auth bridge, wallet, anchors list, pilot, public verify", "src/app/api/knuct/*, register, verify"],
            ["Adapter", "Live HTTP client, mock adapter, circuit breaker, job queue", "src/lib/knuct/*"],
            ["Persistence", "Wallets, anchors, credentials, provision/registration requests", "Prisma models"],
            ["Vendor", "webwallet.knuct.com SAPI/CAPI (+ optional mint/publish URLs)", "KNUCT_BASE_URL"],
        ],
    )
    doc.add_paragraph(
        "Flow (logical): Client → Next.js API → knuct adapter → PostgreSQL (SoR) "
        "and optionally Knuct vendor. Async jobs handle wallet provision, chain publish, credential mint."
    )

    # ── 4. Phased process (what was done) ────────────────────────────────
    doc.add_heading("4. Integration Process (How We Built It)", level=1)

    doc.add_heading("4.1 Phase 0 — Vendor diligence & adapter foundation", level=2)
    for b in [
        "Documented vendor APIs and open questions (docs/KNUCT_Vendor_Questionnaire.md).",
        "Defined adapter interfaces (KnuctAdapter) with MockKnuctAdapter for offline/dev.",
        "Env-based kill switches: KNUCT_ENABLED, KNUCT_ANCHOR_ENABLED, timeouts, retries.",
        "Circuit breaker: after repeated vendor failures, fall back to mock for ~60s.",
        "Redis/Upstash store for multi-instance DID sessions and one-time login grants.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4.2 Phase 1 — Wallet + DID (implemented)", level=2)
    for b in [
        "Wallet create: starttempnode → createwallet → fetch privshare → AES-encrypt → KnuctWallet row.",
        "Privshare download for the owning user (GET /api/knuct/privshare).",
        "DID auth: client hashes privshare → challenge → NLSS response → complete → persist DID.",
        "Login bridge: public POST /api/knuct/login issues short-lived loginToken → NextAuth provider id=knuct.",
        "Self-registration: /register + KnuctDIDAuthPanel → pending KnuctRegistrationRequest → admin approve.",
        "Wallet provision requests: user requests create/reprovision; admin approves; queueWalletProvision.",
        "Optional auto-queue on user create (KNUCT_WALLET_ON_USER_CREATE).",
        "Campus identity mode: password_only | hybrid | knuct_based (settings DB + policy-gate).",
        "UI: My Knuct Wallet on Dashboard; Settings → Knuct; Users wallet badges.",
        "Pilot tooling: POST /api/knuct/pilot and npm run knuct:pilot for cohort provisioning.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4.3 Phase 2 — Verifiable credentials (scaffolded, vendor-gated)", level=2)
    for b in [
        "KnuctCredential model + credential-service / credential-client.",
        "Admin UI panel and GET/POST /api/knuct/credentials.",
        "Mint/verify only when KNUCT_CREDENTIALS_ENABLED and mint/verify URLs are set.",
        "Intended types: attendance_certificate, grade_transcript, compliance_report.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4.4 Phase 3 — Anchors + optional chain publish (hash live; chain gated)", level=2)
    for b in [
        "enqueueAnchor writes SHA-256 payloadHash into BlockchainAnchor (PostgreSQL).",
        "Domain hooks: attendance session complete, violation review, geofence create, "
        "calendar publish, LMS grade publish, subject publish, condonation decision, registration approve.",
        "Public verification: /verify and GET /api/verify/anchor.",
        "Optional chain publish job when KNUCT_CHAIN_PUBLISH_ENABLED + KNUCT_CHAIN_PUBLISH_URL are set.",
        "Audit enrichment: anchorHash linked into audit views where applicable.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # ── 5. End-to-end process flows ──────────────────────────────────────
    doc.add_heading("5. End-to-End Process Flows (Operational)", level=1)

    doc.add_heading("5.1 Enable live Knuct (ops checklist)", level=2)
    add_table(
        doc,
        ["Step", "Action", "Verify"],
        [
            ["1", "Set KNUCT_ENABLED=true, KNUCT_BASE_URL, KNUCT_PRIVSHARE_ENC_KEY", ".env / Vercel env"],
            ["2", "Optional: API_KEY / SECRET / TENANT_ID from vendor", "Vendor credentials"],
            ["3", "Set UPSTASH_REDIS_* for multi-instance DID grants", "GET /api/health redis=ok"],
            ["4", "npm run test:knuct:live", "starttempnode → createwallet → privshare"],
            ["5", "Run pilot cohort (UI or npm run knuct:pilot:sync)", "Wallets active in Users"],
            ["6", "Monitor Dashboard Knuct Ops + /api/health knuct block", "Circuit closed; queue drained"],
            ["7", "Rollback: KNUCT_ENABLED=false", "Core ATMS continues; live provision stops"],
        ],
    )

    doc.add_heading("5.2 DID login process", level=2)
    for b in [
        "User opens /login → Knuct DID tab (if identity mode allows).",
        "Uploads privshare → client computes hash (priv-share.ts / OpenCV on mobile).",
        "POST /api/knuct/login step=challenge → vendor /sapi/auth/challenge.",
        "Client NLSS response → step=complete → vendor auth/response + startnode + walletdata.",
        "Server creates one-time login grant; client signIn('knuct') with loginToken.",
        "NextAuth issues JWT; knuct-persistent-session stores cookies for CAPI; revoke on sign-out.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5.3 Self-registration process", level=2)
    for b in [
        "Public /register (when selfRegistrationEnabled).",
        "DID prove ownership (existing wallet) or create-wallet path → pending request.",
        "Admin Users → Registration Requests → approve → User + KnuctWallet + notification + optional anchor.",
        "Reject stores reason; user notified.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5.4 Wallet provision process", level=2)
    for b in [
        "Admin Provision from Users, or user submits wallet-request, or auto on create.",
        "Job queue calls wallet-service → live or mock adapter.",
        "Success: status=active + DID + encrypted privshare; Failure: status=failed + lastError.",
        "User downloads privshare once stored; used for future DID login.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5.5 Anchor process", level=2)
    for b in [
        "Domain API mutates resource (e.g. session completed).",
        "enqueueAnchor(resourceType, id, payload) → hash stored in BlockchainAnchor.",
        "If chain publish enabled: job POSTs hash to vendor publish URL → knuctTxRef.",
        "Anyone can verify at /verify with hash (or prefix).",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # ── 6. What is live vs gated today ───────────────────────────────────
    doc.add_heading("6. Current Status Matrix", level=1)
    add_table(
        doc,
        ["Capability", "Status in ATMS", "Depends on"],
        [
            ["Mock wallet / DID (dev)", "Complete", "Nothing (default when live off)"],
            ["Live wallet provision", "Complete (pilot-ready)", "KNUCT_ENABLED + vendor sandbox/prod"],
            ["DID login + NextAuth bridge", "Complete", "Redis recommended on Vercel"],
            ["Self-registration + admin approve", "Complete", "auth.self_registration + Knuct"],
            ["Identity mode policy", "Complete", "Settings DB"],
            ["Privshare encrypt/download", "Complete", "KNUCT_PRIVSHARE_ENC_KEY"],
            ["Hash anchors in PostgreSQL", "Complete", "KNUCT_ANCHOR_ENABLED / settings"],
            ["Public /verify", "Complete", "Anchors written"],
            ["Chain publish to Knuct", "Code ready, gated", "Vendor publish URL + enable flag"],
            ["Credential mint/verify", "Code ready, gated", "Vendor mint/verify APIs + enable flag"],
            ["CAPI balances UI", "Complete when session alive", "Successful DID login session"],
            ["Campus-wide production SLA", "Not yet", "Vendor Phase 0 answers + staging"],
        ],
    )

    # ── 7. FUTURE REQUIREMENTS ───────────────────────────────────────────
    doc.add_heading("7. What Is Required in Future", level=1)
    doc.add_paragraph(
        "The following items are required to move from pilot-capable integration to "
        "production campus rollout. Grouped by owner."
    )

    doc.add_heading("7.1 From Knuct vendor (blocking / high priority)", level=2)
    add_table(
        doc,
        ["Requirement", "Why needed", "Unblocks"],
        [
            ["Written answers to vendor questionnaire (Q1–Q10)", "Privshare format, errors, tenancy, DPDP", "Phase 0 sign-off"],
            ["Dedicated staging tenant + credentials", "Do not use shared public sandbox for prod traffic", "Safe live pilot / UAT"],
            ["Production auth (API key / tenant isolation)", "Multi-institution security", "Prod live mode"],
            ["Documented cert mint + verify APIs + schemas", "Attendance certs & transcripts", "Phase 2 credentials go live"],
            ["Chain publish API (stable URL + contract)", "Move from hash-only to on-chain proof", "Phase 3 chain publish"],
            ["starttempnode scaling guidance", "Batch onboard 1000+ students safely", "Mass wallet provision"],
            ["Rate limits, SLA, cost model", "Budget and capacity planning", "Campus-wide rollout"],
            ["W3C DID / resolver roadmap (or confirm CID-only)", "Interoperability with external verifiers", "External verify portals"],
            ["DPA / data residency (India DPDP)", "Legal compliance for student data", "Production beyond pilot"],
            ["Privshare + passphrase lifecycle clarity", "Correct security storage model", "Hardening wallet UX"],
        ],
    )

    doc.add_heading("7.2 From AIMSCS / ATMS engineering", level=2)
    add_table(
        doc,
        ["Requirement", "Why needed", "Suggested next step"],
        [
            ["Production env hardening", "Strong PRIVSHARE_ENC_KEY, Redis, secrets rotation", "Vercel prod checklist"],
            ["Bounded live pilot (one dept / cohort)", "Validate vendor before campus-wide", "knuct:pilot + monitor health"],
            ["Runbook + on-call alerts", "Circuit open, failed wallets, queue depth", "Alert on /api/health knuct"],
            ["User education: privshare custody", "Students lose PNG = lose DID login", "In-app + mailer guidance"],
            ["Exam Cell / hall-ticket hard gate (optional)", "Use condonation clearance + anchors in exam flow", "New Exam module"],
            ["External verifier portal polish", "/verify UX for employers/parents without login", "Branding + QR on reports"],
            ["Internal did:web mapping (if vendor non-W3C)", "Avoid vendor lock-in on CID-only IDs", "Mapping table + resolver"],
            ["Load test wallet provision", "Confirm concurrency before semester start", "Scripted batch + metrics"],
            ["Keep KNUCT_SCMS_Integration_Report.md current", "Older sections say 'no blockchain' — stale", "Sync with this doc"],
            ["Automated E2E in CI (mock + optional live nightly)", "Regressions on login/anchor paths", "Extend test:knuct:*"],
        ],
    )

    doc.add_heading("7.3 Campus / process requirements", level=2)
    for b in [
        "Approve identity mode policy (password_only vs hybrid vs knuct_based) per semester.",
        "Define who approves registration and wallet provision requests (Admin vs HOD).",
        "Decide whether wallet is mandatory for students before exams / hall tickets.",
        "Communicate that ATMS remains usable if Knuct is disabled (attendance % still in Postgres).",
        "Retain audit + anchor hashes for institutional compliance periods.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("7.4 Recommended sequencing (future)", level=2)
    add_table(
        doc,
        ["Order", "Milestone", "Exit criteria"],
        [
            ["1", "Close Phase 0 with vendor written answers", "Questionnaire complete; staging tenant issued"],
            ["2", "Staging live pilot (≤25 users)", "test:knuct:live green; wallets active; no circuit storms"],
            ["3", "Enable chain publish on staging", "Anchors show knuctTxRef; /verify shows chainPublished"],
            ["4", "Credential mint PoC (1 certificate type)", "Mint + public verify without SCMS login"],
            ["5", "Department-wide rollout", "Runbook signed; Redis HA; support trained"],
            ["6", "Campus-wide + exam eligibility hooks", "Policy approved; SLA met for peak enrollment"],
        ],
    )

    # ── 8. Risks ─────────────────────────────────────────────────────────
    doc.add_heading("8. Risks & Mitigations", level=1)
    add_table(
        doc,
        ["Risk", "Mitigation in ATMS today", "Still needed"],
        [
            ["Vendor sandbox unstable", "Circuit breaker + mock adapter", "Dedicated staging + SLA"],
            ["Lost privshare", "Encrypted download once; user custody", "Recovery UX / admin re-provision policy"],
            ["Multi-instance login grant miss", "Upstash Redis store", "Ensure Redis on all prod instances"],
            ["Non-W3C DID", "Store CID as string on KnuctWallet", "Interop mapping if external verify required"],
            ["Over-coupling to Knuct", "Adapter + env kill switches", "Keep Postgres SoR forever"],
        ],
    )

    # ── 9. Related documents ─────────────────────────────────────────────
    doc.add_heading("9. Related Documents", level=1)
    add_table(
        doc,
        ["Document", "Contents"],
        [
            ["KNUCT_Integration_Pointers.docx", "File/API inventory of every Knuct touchpoint"],
            ["KNUCT_SCMS_Integration_Report.md", "Original architecture assessment (update stale sections)"],
            ["docs/KNUCT_Pilot_Setup.md", "Ops steps to enable and run a live pilot"],
            ["docs/KNUCT_Vendor_Questionnaire.md", "Phase 0 diligence questions for Knuct"],
            ["docs/FLOWS.md", "Product flows including Knuct wallet/anchors"],
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"\n— End of Knuct Integration Process & Future Requirements —\n"
        f"Generated by scripts/generate-knuct-integration-process.py ({tip})"
    ).font.size = Pt(9)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
