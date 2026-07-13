"""Generate Knuct Integration Pointers inventory as a Word document."""
from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "KNUCT_Integration_Pointers.docx"
NAVY = "1A3C6E"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = NAVY,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if c_idx == 0:
                        run.font.name = "Consolas"
    doc.add_paragraph()


def git_tip() -> str:
    try:
        return (
            subprocess.check_output(
                ["git", "rev-parse", "--short", "HEAD"],
                cwd=ROOT,
                stderr=subprocess.DEVNULL,
                text=True,
            ).strip()
            or "unknown"
        )
    except Exception:
        return "unknown"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    tip = git_tip()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ── Title ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Knuct Integration Pointers")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "AIMSCS — Academic Tracking & Management System (ATMS)"
    ).font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Complete inventory of Knuct wiring in the codebase\n"
        f"Report date: {date.today().strftime('%d %B %Y')}  ·  Git tip: {tip}"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # ── 1. Executive summary ─────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Knuct is integrated into ATMS as the campus decentralized-identity (DID) and "
        "tamper-evident audit layer. Core capabilities in production code:"
    )
    for bullet in [
        "DID login and self-registration (privshare + NLSS challenge/response).",
        "Per-user Knuct wallets with encrypted privshare storage and download.",
        "SHA-256 blockchain anchors for attendance, violations, geofences, calendar, "
        "LMS grades, subject publish, and condonation decisions (optional chain publish).",
        "Verifiable credentials issuance (gated on vendor mint/verify URLs).",
        "Pilot cohort wallet provisioning for admins.",
        "Campus identity mode policy: password_only | hybrid | knuct_based.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_paragraph(
        "When KNUCT_ENABLED is false or the circuit breaker is open, MockKnuctAdapter "
        "keeps UI flows working without the live vendor."
    )

    # ── 2. UI navigation map ─────────────────────────────────────────────
    doc.add_heading("2. UI Navigation Map", level=1)
    add_table(
        doc,
        ["UI entry", "Knuct features"],
        [
            ["/login", "Knuct DID tab (hybrid / knuct_based identity modes)"],
            ["/register", "Self-registration with KnuctDIDAuthPanel"],
            ["/verify", "Public anchor hash verification"],
            ["Dashboard", "My Knuct Wallet panel (all users); Knuct Ops / pilot (super_admin)"],
            ["Settings → Knuct", "Wallet, DID auth, anchors, credentials, pilot, verify link"],
            ["Settings → Users", "Registration requests + wallet provision request panels"],
            ["Settings → General / Runtime", "Knuct campus details; live / anchors / chain flags"],
            ["Users section", "Wallet badges, per-user Knuct details, admin provision"],
        ],
    )

    # ── 3. API route index ───────────────────────────────────────────────
    doc.add_heading("3. API Route Index", level=1)
    add_table(
        doc,
        ["Route", "Methods", "Category", "Purpose"],
        [
            ["/api/knuct", "GET, POST", "Wallet", "Status, health, stats, queue provision"],
            ["/api/knuct/login", "POST", "Auth", "Public DID login → loginToken"],
            ["/api/knuct/did-auth", "POST", "DID", "Session-bound DID verify/link"],
            ["/api/knuct/account", "GET", "DID/Wallet", "CAPI account + dashboard"],
            ["/api/knuct/privshare", "GET", "Privshare", "Download decrypted privshare PNG"],
            ["/api/knuct/wallet-requests", "GET, POST", "Wallet", "Provision request workflow"],
            ["/api/knuct/anchors", "GET", "Anchors", "List BlockchainAnchor (super_admin)"],
            ["/api/knuct/credentials", "GET, POST", "Credentials", "List / issue credentials"],
            ["/api/knuct/pilot", "GET, POST", "Pilot", "Pilot status + batch provision"],
            ["/api/register", "POST", "Registration", "Public Knuct self-registration"],
            ["/api/register/requests", "GET, POST", "Registration", "Admin approve/reject"],
            ["/api/auth/methods", "GET", "Auth", "Login capability flags (knuctDid, identityMode)"],
            ["/api/verify/anchor", "GET", "Verify", "Public hash lookup"],
            ["/api/health", "GET", "Ops", "Knuct health, queue, circuit, Redis"],
            ["/api/dashboard", "GET", "Ops", "Includes knuct stats block"],
            ["/api/users", "GET, POST", "Wallet", "Wallet on list; maybeProvisionWalletOnCreate"],
            ["/api/audit", "GET", "Anchors", "Audit logs enriched with anchor hashes"],
            ["/api/attendance/sessions/[id]", "PATCH", "Anchors", "Session complete → anchor"],
            ["/api/attendance/violations", "PATCH", "Anchors", "Violation review → anchor"],
            ["/api/geofences", "POST", "Anchors", "Geofence create → anchor"],
            ["/api/calendar", "POST", "Anchors", "Calendar publish → anchor"],
            ["/api/lms/submissions", "PATCH", "Anchors", "Grade publish → anchor"],
            ["/api/masters/subjects/publish", "POST", "Anchors", "Subject publish → anchor"],
            ["/api/settings/config", "GET, PUT", "Config", "Runtime knuct flags"],
            ["/api/settings/platform", "GET", "Config", "Exposes identityMode"],
        ],
    )

    # ── 4. Auth / login ──────────────────────────────────────────────────
    doc.add_heading("4. Auth / Login", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/app/login/page.tsx", "Password + Knuct DID tabs; signIn('knuct')", "/login"],
            ["src/app/api/knuct/login/route.ts", "Challenge → complete → loginToken", "POST /api/knuct/login"],
            ["src/lib/auth.ts", "NextAuth knuct provider; revoke session on sign-out", "signIn('knuct')"],
            ["src/lib/knuct/login-grant.ts", "One-time Redis/memory tokens bridging DID → NextAuth", "login API"],
            ["src/app/api/auth/methods/route.ts", "Public knuctDid / identityMode / preferKnuctLogin", "GET /api/auth/methods"],
            ["src/lib/settings/identity-mode.ts", "Client helpers for Knuct UI visibility", "login + UI"],
            ["src/lib/settings/identity-mode-server.ts", "Server loader for auth.identity_mode", "APIs"],
            ["src/hooks/use-identity-mode.ts", "React hook; knuctUiEnabled", "UI"],
            ["src/middleware.ts", "Bypasses auth for /api/knuct/login", "middleware"],
            ["src/lib/knuct/policy-gate.ts", "403 when identity mode is password_only", "mutating Knuct APIs"],
            ["src/lib/knuct/knuct-persistent-session.ts", "Persist Knuct cookies; revoke on logout", "after DID login"],
            ["src/lib/email.ts", "Welcome email mentions Knuct DID sign-in", "email"],
        ],
    )

    # ── 5. Registration ──────────────────────────────────────────────────
    doc.add_heading("5. Registration", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/app/register/page.tsx", "Self-registration + KnuctDIDAuthPanel", "/register"],
            ["src/app/api/register/route.ts", "DID challenge/complete + pending request", "POST /api/register"],
            ["src/app/api/register/requests/route.ts", "Admin list/approve/reject", "GET/POST /api/register/requests"],
            ["src/lib/knuct/registration-service.ts", "Validate DID; create user + wallet on approve", "register APIs"],
            ["src/components/users/registration-requests-panel.tsx", "Admin approval UI", "Users"],
            ["src/components/knuct/did-auth-panel.tsx", "mode=register posts to /api/register", "/register"],
            ["prisma/schema.prisma", "KnuctRegistrationRequest model", "DB"],
        ],
    )

    # ── 6. Wallet ────────────────────────────────────────────────────────
    doc.add_heading("6. Wallet / Provision Requests", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/lib/knuct/wallet-service.ts", "createKnuctWalletBundle, provisionWallet, queue", "lib / job queue"],
            ["src/lib/knuct/wallet-provision-request-service.ts", "User request + admin approve/reject", "/api/knuct/wallet-requests"],
            ["src/app/api/knuct/route.ts", "GET status; POST queue provision", "/api/knuct"],
            ["src/app/api/knuct/wallet-requests/route.ts", "List/submit/approve provision requests", "wallet-requests API"],
            ["src/app/api/users/route.ts", "maybeProvisionWalletOnCreate; list includes knuctWallet", "/api/users"],
            ["src/components/knuct/my-knuct-wallet-panel.tsx", "User wallet status, privshare, DID, CAPI", "Dashboard"],
            ["src/components/users/wallet-provision-requests-panel.tsx", "Admin review UI", "Users"],
            ["src/components/sections/users-section.tsx", "Wallet badges + admin provision", "Users section"],
            ["src/lib/knuct/stats.ts", "getUserKnuctWallet, dashboard stats", "/api/knuct, dashboard"],
            ["src/lib/knuct/mock-adapter.ts", "Mock wallet/DID when live off or circuit open", "dev/CI"],
            ["prisma/schema.prisma", "KnuctWallet, KnuctWalletProvisionRequest", "DB"],
        ],
    )

    # ── 7. DID / privshare ───────────────────────────────────────────────
    doc.add_heading("7. DID / Privshare / NLSS", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/components/knuct/did-auth-panel.tsx", "Privshare upload → challenge → NLSS → complete", "login/register/settings"],
            ["src/app/api/knuct/did-auth/route.ts", "Authenticated DID verify/link", "POST /api/knuct/did-auth"],
            ["src/lib/knuct/did-auth-flow.ts", "Challenge/complete orchestration", "lib"],
            ["src/lib/knuct/did-auth-session.ts", "In-flight DID cookie jar store", "lib"],
            ["src/lib/knuct/nlss.ts", "NLSS challenge-response (client)", "did-auth-panel"],
            ["src/lib/knuct/priv-share.ts", "Client privshare image → MD5/hash", "did-auth-panel"],
            ["src/lib/knuct/priv-share-server.ts", "Server privshare hash (E2E tests)", "test:knuct:live"],
            ["src/lib/crypto.ts", "AES-256-GCM encrypt privshare (KNUCT_PRIVSHARE_ENC_KEY)", "wallet + privshare API"],
            ["src/app/api/knuct/privshare/route.ts", "Download decrypted privshare PNG", "GET /api/knuct/privshare"],
            ["src/lib/knuct/device.ts", "Mobile UA detection", "client"],
            ["src/lib/knuct/opencv-loader.ts", "Loads /opencv.js for mobile privshare", "client"],
            ["public/opencv.js", "OpenCV.js bundle for mobile DID auth", "/opencv.js"],
            ["src/lib/knuct/capi-service.ts", "CAPI account/dashboard via DID session", "/api/knuct/account"],
            ["src/lib/knuct/account-view.ts", "Parse CAPI for UI balances", "wallet UI"],
            ["src/app/api/knuct/account/route.ts", "Live CAPI bundle for user", "GET /api/knuct/account"],
        ],
    )

    # ── 8. Anchors ───────────────────────────────────────────────────────
    doc.add_heading("8. Anchors + Chain Publish", level=1)
    doc.add_paragraph(
        "Anchor resource types: attendance_session, violation_review, geofence_policy, "
        "calendar_event, grade_publish, subject_publish, condonation_decision."
    )
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/lib/knuct/anchor-service.ts", "SHA-256 anchors in Postgres; enqueueAnchor", "lib"],
            ["src/lib/knuct/chain-publish.ts", "Optional vendor chain publish → knuctTxRef", "job queue"],
            ["src/lib/knuct/anchor-audit.ts", "Enrich audit logs with anchorHash", "/api/audit"],
            ["src/app/api/knuct/anchors/route.ts", "Super-admin list BlockchainAnchor", "GET /api/knuct/anchors"],
            ["src/app/api/attendance/sessions/[id]/route.ts", "Anchor completed sessions", "API"],
            ["src/app/api/attendance/violations/route.ts", "Anchor violation reviews", "API"],
            ["src/app/api/geofences/route.ts", "Anchor geofence policies", "API"],
            ["src/app/api/calendar/route.ts", "Anchor calendar events", "API"],
            ["src/app/api/lms/submissions/route.ts", "Anchor grade publish", "API"],
            ["src/app/api/masters/subjects/publish/route.ts", "Anchor subject publish", "API"],
            ["src/lib/condonation-service.ts", "Anchor condonation decisions", "service"],
            ["src/lib/knuct/registration-service.ts", "Anchor registration approval", "on approve"],
            ["src/app/verify/page.tsx", "Public verify UI", "/verify"],
            ["src/app/api/verify/anchor/route.ts", "Public hash lookup", "GET /api/verify/anchor"],
            ["prisma/schema.prisma", "BlockchainAnchor model", "DB"],
        ],
    )

    # ── 9. Credentials ───────────────────────────────────────────────────
    doc.add_heading("9. Credentials", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/lib/knuct/credential-service.ts", "Issue credentials (DB + async mint)", "lib"],
            ["src/lib/knuct/credential-client.ts", "mintCredentialOnKnuct / verifyCredentialOnKnuct", "lib"],
            ["src/app/api/knuct/credentials/route.ts", "GET list/stats; POST issue (admin)", "/api/knuct/credentials"],
            ["src/components/knuct/knuct-credentials-panel.tsx", "Admin list/issue UI", "Settings → Knuct"],
            ["prisma/schema.prisma", "KnuctCredential model", "DB"],
        ],
    )

    # ── 10. Pilot ────────────────────────────────────────────────────────
    doc.add_heading("10. Pilot", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/lib/knuct/pilot-service.ts", "Batch wallet provision for cohort", "lib"],
            ["src/app/api/knuct/pilot/route.ts", "Status GET; start POST", "/api/knuct/pilot"],
            ["src/components/sections/dashboard-section.tsx", "KnuctOpsPanel (super_admin)", "Dashboard"],
            ["src/components/sections/settings-section.tsx", "Run live pilot button", "Settings → Knuct"],
            ["scripts/knuct-pilot-cohort.ts", "CLI pilot runner", "npm run knuct:pilot"],
            ["docs/KNUCT_Pilot_Setup.md", "Operational pilot guide", "docs"],
        ],
    )

    # ── 11. Settings / identity ──────────────────────────────────────────
    doc.add_heading("11. Settings / Identity Mode", level=1)
    doc.add_paragraph(
        "Identity modes: password_only (Knuct mutations blocked), hybrid (password primary), "
        "knuct_based (DID primary on login)."
    )
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["src/components/sections/settings-section.tsx", "Full Knuct admin tab", "Settings → Knuct"],
            ["src/components/administration/knuct-campus-details.tsx", "Campus Knuct status card", "Settings → General"],
            ["src/components/administration/settings-workspace.tsx", "Runtime integrations panel", "Settings workspace"],
            ["src/lib/settings/registry.ts", "auth.identity_mode, knuct_anchors, integrations.knuct_*", "settings DB"],
            ["src/lib/settings/adapters/system-config.ts", "Maps policies.knuct_anchors_enabled", "adapter"],
            ["src/lib/system-config.ts", "Runtime knuct.liveEnabled / anchors / chainPublish", "/api/settings/config"],
            ["src/lib/system-config-defaults.ts", "Default knuctAnchorsEnabled: true", "defaults"],
        ],
    )

    # ── 12. Libs ─────────────────────────────────────────────────────────
    doc.add_heading("12. Libraries (src/lib/knuct/)", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["index.ts", "Public barrel exports", "imports"],
            ["config.ts", "getKnuctConfig, isKnuctLiveEnabled", "config"],
            ["types.ts", "Adapter interfaces, seed words, dashboard types", "types"],
            ["knuct-client.ts", "Live HTTP adapter to Knuct SAPI + CAPI", "vendor"],
            ["mock-adapter.ts", "Offline mock adapter", "fallback"],
            ["vendor-http.ts", "Shared auth headers for chain/credential APIs", "vendor"],
            ["wallet-service.ts", "Wallet create/provision/queue", "wallet"],
            ["wallet-provision-request-service.ts", "Approval-gated provision", "wallet"],
            ["registration-service.ts", "Self-registration lifecycle", "register"],
            ["did-auth-flow.ts", "DID challenge/complete", "DID"],
            ["did-auth-session.ts", "In-flight DID session store", "DID"],
            ["login-grant.ts", "One-time NextAuth bridge tokens", "auth"],
            ["knuct-persistent-session.ts", "Per-user Knuct cookie persistence", "auth"],
            ["priv-share.ts / priv-share-server.ts", "Privshare hashing", "DID"],
            ["nlss.ts / device.ts / opencv-loader.ts", "NLSS, mobile, OpenCV", "DID"],
            ["anchor-service.ts / chain-publish.ts / anchor-audit.ts", "Anchors + chain + audit", "anchors"],
            ["credential-service.ts / credential-client.ts", "Credential issue/verify", "credentials"],
            ["capi-service.ts / account-view.ts", "CAPI fetch + UI parse", "wallet UI"],
            ["pilot-service.ts", "Cohort provisioning", "pilot"],
            ["stats.ts", "Health, adapter selection, dashboard stats", "ops"],
            ["job-queue.ts", "Async Knuct jobs", "infra"],
            ["redis-store.ts", "Upstash/memory KV for sessions/grants", "infra"],
            ["circuit-breaker.ts", "Failure threshold → mock fallback", "resilience"],
            ["policy-gate.ts", "Campus identity mode API gate", "policy"],
            ["src/lib/crypto.ts", "Privshare AES at rest (outside knuct/)", "crypto"],
        ],
    )

    # ── 13. Database ─────────────────────────────────────────────────────
    doc.add_heading("13. Database (Prisma)", level=1)
    add_table(
        doc,
        ["Artifact", "Purpose", "Entry"],
        [
            ["KnuctWallet", "DID + encrypted privshare + status", "schema.prisma"],
            ["KnuctRegistrationRequest", "Pending self-registrations", "schema.prisma"],
            ["KnuctWalletProvisionRequest", "Admin-approved create/reprovision", "schema.prisma"],
            ["BlockchainAnchor", "Hash anchors + optional knuctTxRef", "schema.prisma"],
            ["KnuctCredential", "Verifiable credential records", "schema.prisma"],
            ["migrations/20250705200000_knuct_chain_credentials", "Chain columns + credentials table", "migration"],
            ["migrations/20250707120000_knuct_registration", "Registration requests + DID index", "migration"],
            ["migrations/20250707133000_registration_wallet_create", "Registration privshare + walletSource", "migration"],
            ["migrations/20250708130000_knuct_wallet_provision_requests", "Provision request table", "migration"],
        ],
    )

    # ── 14. Env vars ─────────────────────────────────────────────────────
    doc.add_heading("14. Environment Variables", level=1)
    doc.add_paragraph("Documented in .env.example:")
    add_table(
        doc,
        ["Variable", "Purpose", "Entry"],
        [
            ["KNUCT_ENABLED", "Master switch for live adapter", ".env"],
            ["KNUCT_BASE_URL", "Vendor base (default webwallet.knuct.com)", ".env"],
            ["KNUCT_WALLET_ON_USER_CREATE", "Auto-queue wallet on user create", ".env"],
            ["KNUCT_PRIVSHARE_ENC_KEY", "AES key for privshare at rest", ".env"],
            ["KNUCT_MAX_RETRIES", "Wallet provision retries", ".env"],
            ["KNUCT_CIRCUIT_BREAKER_THRESHOLD", "Failures before circuit opens", ".env"],
            ["KNUCT_ANCHOR_ENABLED", "Global anchor kill switch", ".env"],
            ["KNUCT_PILOT_COHORT_LIMIT", "Default pilot batch size", ".env"],
            ["KNUCT_API_KEY / SECRET / TENANT_ID", "Optional vendor auth headers", ".env"],
            ["KNUCT_CHAIN_PUBLISH_ENABLED + URL", "Phase 3 chain publish", ".env"],
            ["KNUCT_CREDENTIALS_ENABLED + MINT/VERIFY URLs", "Phase 2 credentials", ".env"],
            ["KNUCT_REQUEST_TIMEOUT_MS / PRIVSHARE_TIMEOUT_MS", "HTTP timeouts", "knuct-client"],
            ["UPSTASH_REDIS_REST_*", "Multi-instance DID/session/grant store", ".env"],
            ["auth.identity_mode (settings DB)", "Campus Knuct policy", "settings"],
            ["policies.knuct_anchors_enabled (settings DB)", "Campus anchor policy", "settings"],
        ],
    )

    # ── 15. Scripts ──────────────────────────────────────────────────────
    doc.add_heading("15. Scripts & Tooling", level=1)
    add_table(
        doc,
        ["File / command", "Purpose", "Entry"],
        [
            ["npm run test:knuct:live", "Full live vendor contract test", "scripts/test-knuct-live.ts"],
            ["npm run test:knuct:anchor", "Anchor creation smoke test", "scripts/test-knuct-anchor.ts"],
            ["npx tsx scripts/test-knuct-provision.ts", "Provision super_admin wallet", "script"],
            ["npm run knuct:pilot / knuct:pilot:sync", "CLI pilot provisioning", "scripts/knuct-pilot-cohort.ts"],
            ["scripts/check-audit-anchors.ts", "Audit ↔ anchor correlation", "manual"],
            ["scripts/generate-knuct-status-report.py", "Older status Word report", "docs"],
            ["scripts/generate-knuct-report.py", "MD → DOCX for integration report", "docs"],
            ["scripts/generate-knuct-integration-pointers.py", "This inventory Word document", "docs"],
        ],
    )

    # ── 16. Docs in repo ─────────────────────────────────────────────────
    doc.add_heading("16. Related Documentation", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Entry"],
        [
            ["KNUCT_SCMS_Integration_Report.md", "Architecture assessment / phased rollout", "repo root"],
            ["docs/KNUCT_Pilot_Setup.md", "Live pilot enablement guide", "docs"],
            ["docs/KNUCT_Vendor_Questionnaire.md", "Vendor diligence (chain, credentials)", "docs"],
            ["docs/KNUCT_Vendor_Email_Draft.md", "Vendor outreach template", "docs"],
            ["docs/KNUCT_Vendor_Reply_Draft.md", "Vendor reply template", "docs"],
            ["docs/FLOWS.md", "Product flows including Knuct wallet/anchors", "docs"],
            ["README.md", "Mentions anchors in condonation features", "README"],
        ],
    )

    # ── 17. Vendor API surface ───────────────────────────────────────────
    doc.add_heading("17. External Knuct Vendor API Surface", level=1)
    add_table(
        doc,
        ["Endpoint", "Purpose", "Entry"],
        [
            ["GET /sapi/starttempnode", "Wallet creation prep", "wallet-service"],
            ["POST /sapi/createwallet", "Create wallet + privshare URL", "wallet-service"],
            ["GET /sapi/privshare?k=...", "Download privshare PNG", "knuct-client"],
            ["POST /sapi/auth/challenge", "DID auth step 1", "DID login"],
            ["POST /sapi/auth/response", "DID auth step 2", "DID login"],
            ["GET /sapi/startnode", "Start node after auth", "DID login"],
            ["GET /sapi/walletdata", "Retrieve DID after auth", "DID login"],
            ["GET /sapi/logout", "End Knuct session", "logout"],
            ["GET /capi/check", "Session check", "capi-service"],
            ["GET /capi/start", "CAPI node sync", "capi-service"],
            ["GET /capi/getAccountInfo", "Account details", "capi-service"],
            ["GET /capi/getDashboard", "Token/balance dashboard", "capi-service"],
            ["POST {KNUCT_CHAIN_PUBLISH_URL}", "Publish anchor hash to chain", "chain-publish"],
            ["POST {KNUCT_CREDENTIAL_MINT_URL}", "Mint verifiable credential", "credential-client"],
            ["GET {KNUCT_CREDENTIAL_VERIFY_URL}", "Verify credential asset", "credential-client"],
        ],
    )

    # ── 18. Notes ────────────────────────────────────────────────────────
    doc.add_heading("18. Operational Notes", level=1)
    for note in [
        "MockKnuctAdapter is used when KNUCT_ENABLED=false or the circuit breaker is open — UI can run without the live vendor.",
        "Chain publish and credential mint are implemented but gated on vendor URL env vars.",
        "Upstash Redis is recommended on multi-instance Vercel so DID challenge→response and login grants work across lambdas.",
        "middleware.ts allows unauthenticated access to /api/knuct/login, /register*, and /api/verify*.",
        "vercel.json has no Knuct-specific routes; all config is via environment variables.",
    ]:
        doc.add_paragraph(note, style="List Bullet")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        f"\n— End of Knuct Integration Pointers —\nGenerated by scripts/generate-knuct-integration-pointers.py ({tip})"
    ).font.size = Pt(9)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
